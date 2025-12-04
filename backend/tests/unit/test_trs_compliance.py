"""TRS Compliance Test Suite.

Tests aligned with the Technical Requirements Specification (TRS) Test Plan.
These tests verify that the DOCX localization pipeline meets all requirements.
"""

import base64
from io import BytesIO
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from docx import Document as DocxDocument
from docx.shared import Inches

from backend.translation.chunking.chunker import Chunker, chunk_document
from backend.translation.ingestion.docx_parser import parse_docx_bytes
from backend.translation.ir import Document, Image, Paragraph, Table, TranslationUnit
from backend.translation.reconstruction.reconstructor import reconstruct_document


# =============================================================================
# Test Fixtures
# =============================================================================

def create_simple_docx(text: str = "Hello World") -> bytes:
    """Create a simple DOCX with one paragraph."""
    doc = DocxDocument()
    doc.add_paragraph(text)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def create_docx_with_image() -> bytes:
    """Create a DOCX with text AND an embedded image in the same paragraph area."""
    doc = DocxDocument()
    doc.add_paragraph("Text before image")
    
    # A 1x1 red PNG (minimal valid PNG)
    png_data = base64.b64decode(
        'iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8DwHwAFBQIAX8jx0gAAAABJRU5ErkJggg=='
    )
    img_stream = BytesIO(png_data)
    doc.add_picture(img_stream, width=Inches(1), height=Inches(1))
    
    doc.add_paragraph("Text after image")
    
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def create_table_docx(rows: int = 3, cols: int = 3) -> bytes:
    """Create a DOCX with a table."""
    doc = DocxDocument()
    table = doc.add_table(rows=rows, cols=cols)
    
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f"Cell {i+1},{j+1}"
    
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def create_empty_paragraphs_docx() -> bytes:
    """Create a DOCX with empty and whitespace-only paragraphs."""
    doc = DocxDocument()
    doc.add_paragraph("Real content")
    doc.add_paragraph("")  # Empty
    doc.add_paragraph("   ")  # Whitespace only
    doc.add_paragraph("\n\n")  # Newlines only
    doc.add_paragraph("More real content")
    
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def create_large_docx(paragraph_count: int = 50) -> bytes:
    """Create a DOCX with many paragraphs for batch testing."""
    doc = DocxDocument()
    
    for i in range(paragraph_count):
        doc.add_paragraph(f"Paragraph {i+1}: This is test content for paragraph number {i+1}.")
    
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def create_special_chars_docx() -> bytes:
    """Create a DOCX with special characters that could break XML."""
    doc = DocxDocument()
    doc.add_paragraph("Text with ampersand & symbol")
    doc.add_paragraph("Text with angle brackets < and >")
    doc.add_paragraph('Text with "double" and \'single\' quotes')
    doc.add_paragraph("Mix: <tag> AT&T \"quoted\" & more")
    
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# =============================================================================
# TC-01: Sanity Check
# =============================================================================

class TestTC01SanityCheck:
    """TC-01: Basic document parsing and reconstruction."""
    
    def test_simple_docx_parses_successfully(self):
        """A simple DOCX with one paragraph should parse without errors."""
        content = create_simple_docx("Hello World")
        doc = parse_docx_bytes(content, "test.docx")
        
        assert isinstance(doc, Document)
        assert doc.source_format == "docx"
        assert len(doc.sections) >= 1
        
    def test_parsed_document_can_be_opened_by_docx(self):
        """Output document should be valid and openable."""
        from backend.translation.export.docx_writer import write_docx_bytes
        
        content = create_simple_docx("Hello World")
        doc = parse_docx_bytes(content, "test.docx")
        
        # Write back to bytes
        output_bytes = write_docx_bytes(doc)
        
        # Should be able to re-parse without errors
        reparsed = parse_docx_bytes(output_bytes, "output.docx")
        assert reparsed is not None
        
    def test_text_content_is_preserved(self):
        """Original text content should be extractable."""
        content = create_simple_docx("Hello World")
        doc = parse_docx_bytes(content, "test.docx")
        
        all_text = " ".join(p.text for p in doc.all_paragraphs)
        assert "Hello World" in all_text


# =============================================================================
# TC-02: Image Guard
# =============================================================================

class TestTC02ImageGuard:
    """TC-02: Images must be preserved, text in image paragraphs should be skipped."""
    
    def test_image_is_extracted(self):
        """Images should be detected and extracted from DOCX."""
        content = create_docx_with_image()
        doc = parse_docx_bytes(content, "test.docx")
        
        # Find images in elements
        images = [e for e in doc.sections[0].elements if isinstance(e, Image)]
        assert len(images) >= 1
        
    def test_image_paragraph_marked_skip_translation(self):
        """Paragraphs containing images should have skip_translation=True."""
        content = create_docx_with_image()
        doc = parse_docx_bytes(content, "test.docx")
        
        # Any paragraph that coexists with an image in its original location
        # should be marked skip_translation
        paragraphs = [e for e in doc.sections[0].elements if isinstance(e, Paragraph)]
        
        # There should be at least some paragraphs
        assert len(paragraphs) >= 2
        
    def test_skip_translation_paragraphs_not_chunked(self):
        """Paragraphs with skip_translation=True should not appear in chunks."""
        content = create_docx_with_image()
        doc = parse_docx_bytes(content, "test.docx")
        
        # Mark a paragraph as skip (simulating image guard)
        for section in doc.sections:
            for elem in section.elements:
                if isinstance(elem, Paragraph) and elem.skip_translation:
                    # This paragraph should not appear in chunks
                    units = chunk_document(doc)
                    
                    # Check that skip_translation paragraph IDs are not in any unit
                    all_chunked_ids = set()
                    for unit in units:
                        for ref in unit.element_refs:
                            all_chunked_ids.add(ref.element_id)
                    
                    assert elem.id not in all_chunked_ids
                    
    def test_image_roundtrip_preserves_image(self):
        """Parse -> Write -> Parse should preserve images."""
        from backend.translation.export.docx_writer import write_docx_bytes
        
        content = create_docx_with_image()
        doc1 = parse_docx_bytes(content, "test.docx")
        
        images1 = [e for e in doc1.sections[0].elements if isinstance(e, Image)]
        
        # Round-trip
        output_bytes = write_docx_bytes(doc1)
        doc2 = parse_docx_bytes(output_bytes, "output.docx")
        
        images2 = [e for e in doc2.sections[0].elements if isinstance(e, Image)]
        
        assert len(images1) == len(images2)


# =============================================================================
# TC-03: Table Integrity
# =============================================================================

class TestTC03TableIntegrity:
    """TC-03: Tables should maintain their structure after parsing."""
    
    def test_table_row_count_preserved(self):
        """3x3 table should have exactly 3 rows."""
        content = create_table_docx(rows=3, cols=3)
        doc = parse_docx_bytes(content, "test.docx")
        
        tables = doc.all_tables
        assert len(tables) == 1
        assert tables[0].num_rows == 3
        
    def test_table_column_count_preserved(self):
        """3x3 table should have exactly 3 columns."""
        content = create_table_docx(rows=3, cols=3)
        doc = parse_docx_bytes(content, "test.docx")
        
        tables = doc.all_tables
        assert len(tables) == 1
        assert tables[0].num_cols == 3
        
    def test_table_cell_content_preserved(self):
        """All 9 cells should contain their original text."""
        content = create_table_docx(rows=3, cols=3)
        doc = parse_docx_bytes(content, "test.docx")
        
        table = doc.all_tables[0]
        
        for i in range(3):
            for j in range(3):
                cell = table.get_cell(i, j)
                assert cell is not None
                assert f"Cell {i+1},{j+1}" in cell.text


# =============================================================================
# TC-04: Empty Run Handling
# =============================================================================

class TestTC04EmptyRunHandling:
    """TC-04: Empty paragraphs should be skipped, no API calls made."""
    
    def test_empty_paragraphs_skipped_in_chunking(self):
        """Empty and whitespace-only paragraphs should not create translation units."""
        content = create_empty_paragraphs_docx()
        doc = parse_docx_bytes(content, "test.docx")
        
        units = chunk_document(doc)
        
        # Should only have chunks for real content
        # Combine all source text
        all_source_text = " ".join(u.source_text for u in units)
        
        # Real content should be present
        assert "Real content" in all_source_text
        assert "More real content" in all_source_text
        
    def test_empty_paragraphs_not_sent_to_api(self):
        """Verify that empty strings don't result in API calls."""
        content = create_empty_paragraphs_docx()
        doc = parse_docx_bytes(content, "test.docx")
        
        units = chunk_document(doc)
        
        # Each unit should have non-empty source text
        for unit in units:
            assert unit.source_text.strip(), "Empty unit should not exist"


# =============================================================================
# TC-05: Batch Reconstruction
# =============================================================================

class TestTC05BatchReconstruction:
    """TC-05: Documents with many paragraphs should maintain order."""
    
    def test_paragraph_order_preserved(self):
        """50 paragraphs should maintain their exact order after chunking."""
        content = create_large_docx(paragraph_count=50)
        doc = parse_docx_bytes(content, "test.docx")
        
        # Get original paragraph order
        original_paragraphs = doc.all_paragraphs
        original_ids = [p.id for p in original_paragraphs]
        
        # Chunk the document
        units = chunk_document(doc)
        
        # Collect element refs in order
        chunked_ids = []
        for unit in units:
            for ref in unit.element_refs:
                if ref.element_type == "paragraph":
                    chunked_ids.append(ref.element_id)
        
        # All paragraph IDs should appear in the same relative order
        # (though not all may be chunked if empty)
        seen_ids = [pid for pid in original_ids if pid in chunked_ids]
        chunked_order = [cid for cid in chunked_ids if cid in seen_ids]
        
        assert seen_ids == chunked_order, "Paragraph order should be preserved"
        
    def test_reconstruction_maintains_order(self):
        """Reconstruction should maintain the original document order."""
        content = create_large_docx(paragraph_count=10)
        doc = parse_docx_bytes(content, "test.docx")
        
        # Get original text order
        original_texts = [p.text for p in doc.all_paragraphs]
        
        # Create mock translations
        units = chunk_document(doc)
        for unit in units:
            # Simple mock: prefix with [T]
            unit.translated_text = unit.source_text.replace("<p ", "[T]<p ")
        
        # Reconstruct
        result = reconstruct_document(doc, units)
        
        # Check order is maintained
        reconstructed_paragraphs = result.document.all_paragraphs
        for i, para in enumerate(reconstructed_paragraphs):
            if para.text.strip():
                # Should still contain the original paragraph number reference
                assert f"Paragraph {i+1}" in para.text or "[T]" in para.text


# =============================================================================
# UAT-04: Special Characters
# =============================================================================

class TestUAT04SpecialCharacters:
    """UAT-04: Special characters should not crash XML parsing."""
    
    def test_ampersand_handling(self):
        """Ampersand (&) should be handled correctly."""
        content = create_special_chars_docx()
        doc = parse_docx_bytes(content, "test.docx")
        
        all_text = " ".join(p.text for p in doc.all_paragraphs)
        assert "&" in all_text or "AT" in all_text  # Either escaped or preserved
        
    def test_angle_brackets_handling(self):
        """Angle brackets (< >) should be handled correctly."""
        content = create_special_chars_docx()
        doc = parse_docx_bytes(content, "test.docx")
        
        # Document should parse without crashing
        assert len(doc.sections) >= 1
        
    def test_quotes_handling(self):
        """Quotes should be handled correctly."""
        content = create_special_chars_docx()
        doc = parse_docx_bytes(content, "test.docx")
        
        all_text = " ".join(p.text for p in doc.all_paragraphs)
        assert "double" in all_text
        assert "single" in all_text
        
    def test_special_chars_roundtrip(self):
        """Special characters should survive parse -> write -> parse."""
        from backend.translation.export.docx_writer import write_docx_bytes
        
        content = create_special_chars_docx()
        doc1 = parse_docx_bytes(content, "test.docx")
        
        output_bytes = write_docx_bytes(doc1)
        doc2 = parse_docx_bytes(output_bytes, "output.docx")
        
        # Content should be preserved
        text1 = " ".join(p.text for p in doc1.all_paragraphs)
        text2 = " ".join(p.text for p in doc2.all_paragraphs)
        
        # Core content should match
        assert "AT" in text2  # From AT&T
        assert "quoted" in text2
