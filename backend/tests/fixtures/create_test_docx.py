"""Script to create test DOCX files for testing the parser."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def create_simple_docx(output_path: Path) -> None:
    """Create a simple test DOCX with basic content."""
    doc = Document()

    # Add a simple paragraph
    doc.add_paragraph("Hello, World! 안녕하세요!")

    # Add paragraph with multiple runs
    para2 = doc.add_paragraph()
    para2.add_run("This is ")
    bold_run = para2.add_run("bold")
    bold_run.bold = True
    para2.add_run(" and ")
    italic_run = para2.add_run("italic")
    italic_run.italic = True
    para2.add_run(" text.")

    doc.save(str(output_path))


def create_styled_docx(output_path: Path) -> None:
    """Create a DOCX with various styles and formatting."""
    doc = Document()

    # Add heading
    doc.add_heading("Document Title", level=1)

    # Add styled paragraph
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)

    run = para.add_run("Centered text with custom spacing")
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(255, 0, 0)  # Red

    # Add normal paragraph
    doc.add_paragraph("This is a normal paragraph with default styling.")

    doc.save(str(output_path))


def create_table_docx(output_path: Path) -> None:
    """Create a DOCX with a table."""
    doc = Document()

    doc.add_paragraph("Here is a table:")

    # Create a 3x3 table
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"

    # Fill the table
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f"Cell {i+1},{j+1}"

    # Add Korean content in a cell
    table.rows[1].cells[1].paragraphs[0].clear()
    table.rows[1].cells[1].paragraphs[0].add_run("한국어 텍스트")

    doc.add_paragraph("Table ends here.")

    doc.save(str(output_path))


def create_complex_docx(output_path: Path) -> None:
    """Create a complex DOCX with mixed content."""
    doc = Document()

    # Title
    doc.add_heading("테스트 문서 (Test Document)", level=0)

    # Introduction
    intro = doc.add_paragraph()
    intro.add_run("이 문서는 번역 테스트를 위한 샘플입니다. ")
    intro.add_run("(This document is a sample for translation testing.)")

    # Section 1
    doc.add_heading("제1장: 서론 (Chapter 1: Introduction)", level=1)

    para1 = doc.add_paragraph()
    para1.add_run("문서 번역 시스템은 ")
    bold = para1.add_run("정확한 형식 보존")
    bold.bold = True
    para1.add_run("이 중요합니다.")

    # Table section
    doc.add_heading("데이터 테이블 (Data Table)", level=2)

    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"

    table.rows[0].cells[0].text = "항목 (Item)"
    table.rows[0].cells[1].text = "값 (Value)"
    table.rows[1].cells[0].text = "테스트"
    table.rows[1].cells[1].text = "성공"

    # Conclusion
    doc.add_heading("결론 (Conclusion)", level=1)
    doc.add_paragraph("번역이 완료되었습니다. Translation completed.")

    doc.save(str(output_path))


if __name__ == "__main__":
    fixtures_dir = Path(__file__).parent

    create_simple_docx(fixtures_dir / "simple.docx")
    create_styled_docx(fixtures_dir / "styled.docx")
    create_table_docx(fixtures_dir / "table.docx")
    create_complex_docx(fixtures_dir / "complex.docx")

    print("Test DOCX files created successfully!")

