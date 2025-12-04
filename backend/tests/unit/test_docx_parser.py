"""Tests for DOCX parser."""

from io import BytesIO

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from backend.translation.ingestion.docx_parser import parse_docx_bytes
from backend.translation.ir import Alignment, Document, Paragraph, Table


def create_simple_docx_bytes() -> bytes:
    """Create a simple DOCX in memory and return as bytes."""
    doc = DocxDocument()
    doc.add_paragraph("Hello, World!")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def create_styled_docx_bytes() -> bytes:
    """Create a DOCX with styled content."""
    doc = DocxDocument()

    para = doc.add_paragraph()
    para.add_run("Normal text, ")
    bold_run = para.add_run("bold text")
    bold_run.bold = True
    para.add_run(", ")
    italic_run = para.add_run("italic text")
    italic_run.italic = True

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def create_table_docx_bytes() -> bytes:
    """Create a DOCX with a table."""
    doc = DocxDocument()

    doc.add_paragraph("Before table")

    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "A1"
    table.rows[0].cells[1].text = "B1"
    table.rows[1].cells[0].text = "A2"
    table.rows[1].cells[1].text = "B2"

    doc.add_paragraph("After table")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def create_korean_docx_bytes() -> bytes:
    """Create a DOCX with Korean content."""
    doc = DocxDocument()

    doc.add_paragraph("안녕하세요, 세계!")
    doc.add_paragraph("이것은 한국어 테스트입니다.")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


class TestDocxParser:
    """Tests for DocxParser class."""

    def test_parse_simple_docx(self):
        """Test parsing a simple document."""
        content = create_simple_docx_bytes()
        doc = parse_docx_bytes(content, "test.docx")

        assert isinstance(doc, Document)
        assert doc.source_filename == "test.docx"
        assert doc.source_format == "docx"
        assert len(doc.sections) >= 1

    def test_parse_extracts_text(self):
        """Test that text content is extracted."""
        content = create_simple_docx_bytes()
        doc = parse_docx_bytes(content)

        paragraphs = doc.all_paragraphs
        assert len(paragraphs) >= 1

        all_text = " ".join(p.text for p in paragraphs)
        assert "Hello, World!" in all_text

    def test_parse_styled_text(self):
        """Test parsing styled text with bold and italic."""
        content = create_styled_docx_bytes()
        doc = parse_docx_bytes(content)

        paragraphs = doc.all_paragraphs
        assert len(paragraphs) >= 1

        para = paragraphs[0]
        assert len(para.runs) >= 3

        # Find bold run
        bold_runs = [r for r in para.runs if r.bold]
        assert len(bold_runs) >= 1
        assert "bold" in bold_runs[0].text

        # Find italic run
        italic_runs = [r for r in para.runs if r.italic]
        assert len(italic_runs) >= 1
        assert "italic" in italic_runs[0].text

    def test_parse_table(self):
        """Test parsing tables."""
        content = create_table_docx_bytes()
        doc = parse_docx_bytes(content)

        tables = doc.all_tables
        assert len(tables) >= 1

        table = tables[0]
        assert table.num_rows == 2
        assert table.num_cols == 2

        # Check cell content
        assert table.get_cell(0, 0).text == "A1"
        assert table.get_cell(0, 1).text == "B1"
        assert table.get_cell(1, 0).text == "A2"
        assert table.get_cell(1, 1).text == "B2"

    def test_parse_korean_text(self):
        """Test parsing Korean text."""
        content = create_korean_docx_bytes()
        doc = parse_docx_bytes(content)

        paragraphs = doc.all_paragraphs
        assert len(paragraphs) >= 2

        all_text = " ".join(p.text for p in paragraphs)
        assert "안녕하세요" in all_text
        assert "한국어" in all_text

    def test_parse_preserves_paragraph_order(self):
        """Test that paragraph order is preserved."""
        content = create_table_docx_bytes()
        doc = parse_docx_bytes(content)

        # Should have: paragraph, table, paragraph
        section = doc.sections[0]
        assert len(section.elements) == 3
        assert isinstance(section.elements[0], Paragraph)
        assert isinstance(section.elements[1], Table)
        assert isinstance(section.elements[2], Paragraph)

        assert section.elements[0].text == "Before table"
        assert section.elements[2].text == "After table"

    def test_document_serialization_roundtrip(self):
        """Test that parsed document can be serialized and restored."""
        content = create_styled_docx_bytes()
        doc = parse_docx_bytes(content)

        # Serialize to JSON
        json_str = doc.to_json()

        # Restore from JSON
        restored = Document.from_json(json_str)

        # Verify content matches
        assert len(restored.sections) == len(doc.sections)
        assert len(restored.all_paragraphs) == len(doc.all_paragraphs)

        for orig_para, rest_para in zip(doc.all_paragraphs, restored.all_paragraphs):
            assert orig_para.text == rest_para.text


class TestDocxParserWithFontAndColor:
    """Tests for font and color extraction."""

    def test_parse_font_properties(self):
        """Test parsing font name and size."""
        doc = DocxDocument()
        para = doc.add_paragraph()
        run = para.add_run("Styled text")
        run.font.name = "Arial"
        run.font.size = Pt(14)

        buffer = BytesIO()
        doc.save(buffer)
        content = buffer.getvalue()

        parsed = parse_docx_bytes(content)
        paragraphs = parsed.all_paragraphs
        assert len(paragraphs) >= 1

        runs = paragraphs[0].runs
        assert len(runs) >= 1

        # Font name should be preserved
        assert runs[0].font_name == "Arial"
        # Font size should be 14 points
        assert runs[0].font_size == 14.0

    def test_parse_font_color(self):
        """Test parsing font color."""
        doc = DocxDocument()
        para = doc.add_paragraph()
        run = para.add_run("Red text")
        run.font.color.rgb = RGBColor(255, 0, 0)

        buffer = BytesIO()
        doc.save(buffer)
        content = buffer.getvalue()

        parsed = parse_docx_bytes(content)
        runs = parsed.all_paragraphs[0].runs

        assert runs[0].color == "FF0000"


class TestDocxParserWithAlignment:
    """Tests for paragraph alignment extraction."""

    def test_parse_center_alignment(self):
        """Test parsing centered paragraph."""
        doc = DocxDocument()
        para = doc.add_paragraph("Centered text")
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        buffer = BytesIO()
        doc.save(buffer)
        content = buffer.getvalue()

        parsed = parse_docx_bytes(content)
        paragraphs = parsed.all_paragraphs

        assert paragraphs[0].alignment == Alignment.CENTER

    def test_parse_right_alignment(self):
        """Test parsing right-aligned paragraph."""
        doc = DocxDocument()
        para = doc.add_paragraph("Right-aligned text")
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        buffer = BytesIO()
        doc.save(buffer)
        content = buffer.getvalue()

        parsed = parse_docx_bytes(content)
        paragraphs = parsed.all_paragraphs

        assert paragraphs[0].alignment == Alignment.RIGHT


class TestDocxParserWithSpacing:
    """Tests for paragraph spacing extraction."""

    def test_parse_paragraph_spacing(self):
        """Test parsing paragraph spacing."""
        doc = DocxDocument()
        para = doc.add_paragraph("Spaced paragraph")
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)

        buffer = BytesIO()
        doc.save(buffer)
        content = buffer.getvalue()

        parsed = parse_docx_bytes(content)
        paragraphs = parsed.all_paragraphs

        assert paragraphs[0].space_before == 12.0
        assert paragraphs[0].space_after == 6.0


class TestDocxParserWithImages:
    """Tests for image extraction from DOCX files."""
    
    def _create_docx_with_image(self) -> bytes:
        """Create a DOCX with an inline image."""
        import base64
        from docx.shared import Inches
        
        doc = DocxDocument()
        doc.add_paragraph("Text before image")
        
        # A 1x1 red PNG (minimal valid PNG)
        png_data = base64.b64decode(
            'iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8DwHwAFBQIAX8jx0gAAAABJRU5ErkJggg=='
        )
        img_stream = BytesIO(png_data)
        doc.add_picture(img_stream, width=Inches(2), height=Inches(1))
        
        doc.add_paragraph("Text after image")
        
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    
    def test_parse_inline_image(self):
        """Test that inline images are parsed."""
        from backend.translation.ir import Image
        
        content = self._create_docx_with_image()
        doc = parse_docx_bytes(content, "test_image.docx")
        
        # Find images in elements
        images = [e for e in doc.sections[0].elements if isinstance(e, Image)]
        assert len(images) == 1
        
    def test_image_dimensions_preserved(self):
        """Test that image dimensions are extracted correctly."""
        from backend.translation.ir import Image
        
        content = self._create_docx_with_image()
        doc = parse_docx_bytes(content, "test_image.docx")
        
        images = [e for e in doc.sections[0].elements if isinstance(e, Image)]
        assert len(images) == 1
        
        img = images[0]
        assert img.width == 2.0
        assert img.height == 1.0
        assert img.format == "png"
        
    def test_image_data_is_base64(self):
        """Test that image data is stored as valid base64."""
        import base64
        from backend.translation.ir import Image
        
        content = self._create_docx_with_image()
        doc = parse_docx_bytes(content, "test_image.docx")
        
        images = [e for e in doc.sections[0].elements if isinstance(e, Image)]
        img = images[0]
        
        # Should be able to decode the base64 data
        decoded = base64.b64decode(img.data)
        assert len(decoded) > 0
        
    def test_image_roundtrip(self):
        """Test that images survive parse -> write -> parse cycle."""
        from backend.translation.ir import Image
        from backend.translation.export.docx_writer import write_docx_bytes
        
        # Parse original
        content = self._create_docx_with_image()
        doc1 = parse_docx_bytes(content, "test_image.docx")
        
        # Write and re-parse
        output_bytes = write_docx_bytes(doc1)
        doc2 = parse_docx_bytes(output_bytes, "output.docx")
        
        # Check images are preserved
        images1 = [e for e in doc1.sections[0].elements if isinstance(e, Image)]
        images2 = [e for e in doc2.sections[0].elements if isinstance(e, Image)]
        
        assert len(images1) == len(images2) == 1
        assert images1[0].width == images2[0].width
        assert images1[0].height == images2[0].height
        assert images1[0].data == images2[0].data
        
    def test_image_serialization(self):
        """Test that images can be serialized to JSON and restored."""
        from backend.translation.ir import Image
        
        content = self._create_docx_with_image()
        doc = parse_docx_bytes(content, "test_image.docx")
        
        # Serialize to JSON
        json_str = doc.to_json()
        
        # Restore and verify
        restored = Document.from_json(json_str)
        
        images = [e for e in restored.sections[0].elements if isinstance(e, Image)]
        assert len(images) == 1
        assert images[0].width == 2.0
        assert images[0].height == 1.0

