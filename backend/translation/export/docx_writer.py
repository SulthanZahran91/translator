"""DOCX writer for exporting IR to Word documents.

Converts the Intermediate Representation back to DOCX format,
preserving all formatting information.
"""

from pathlib import Path
from io import BytesIO

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from backend.translation.ir import (
    Alignment,
    Document,
    DocumentElement,
    Paragraph,
    Section,
    Table,
    TableCell,
    TableRow,
    TextRun,
)


def _hex_to_rgb(hex_color: str | None) -> RGBColor | None:
    """Convert hex color string to RGBColor."""
    if not hex_color or len(hex_color) != 6:
        return None
    try:
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        return RGBColor(r, g, b)
    except ValueError:
        return None


def _get_docx_alignment(alignment: Alignment) -> WD_ALIGN_PARAGRAPH:
    """Convert IR alignment to docx alignment."""
    alignment_map = {
        Alignment.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
        Alignment.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
        Alignment.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
        Alignment.JUSTIFY: WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)


class DocxWriter:
    """Writes IR documents to DOCX format."""
    
    def __init__(self) -> None:
        self._doc: DocxDocument | None = None
    
    def write(self, document: Document, output_path: Path | str) -> Path:
        """Write document to DOCX file.
        
        Args:
            document: The document IR to write
            output_path: Path to save the DOCX file
            
        Returns:
            Path to the written file
        """
        output_path = Path(output_path)
        
        self._doc = DocxDocument()
        self._write_document(document)
        self._doc.save(str(output_path))
        
        return output_path
    
    def write_bytes(self, document: Document) -> bytes:
        """Write document to bytes.
        
        Args:
            document: The document IR to write
            
        Returns:
            DOCX file content as bytes
        """
        self._doc = DocxDocument()
        self._write_document(document)
        
        buffer = BytesIO()
        self._doc.save(buffer)
        return buffer.getvalue()
    
    def _write_document(self, document: Document) -> None:
        """Write the full document structure."""
        for section in document.sections:
            self._write_section(section)
    
    def _write_section(self, section: Section) -> None:
        """Write a section's elements."""
        if not self._doc:
            return
        
        # Set page dimensions if specified
        # Note: values should already be in inches from the parser
        # but we validate to avoid overflow errors
        if section.page_width or section.page_height:
            docx_section = self._doc.sections[-1]
            
            # Validate reasonable bounds (max ~50 inches for any dimension)
            def safe_inches(value: float | None) -> Inches | None:
                if value is None:
                    return None
                # If value is unreasonably large, it's likely in EMUs - skip
                if value > 100:
                    return None
                return Inches(value)
            
            if section.page_width:
                val = safe_inches(section.page_width)
                if val:
                    docx_section.page_width = val
            if section.page_height:
                val = safe_inches(section.page_height)
                if val:
                    docx_section.page_height = val
            if section.margin_top is not None:
                val = safe_inches(section.margin_top)
                if val:
                    docx_section.top_margin = val
            if section.margin_bottom is not None:
                val = safe_inches(section.margin_bottom)
                if val:
                    docx_section.bottom_margin = val
            if section.margin_left is not None:
                val = safe_inches(section.margin_left)
                if val:
                    docx_section.left_margin = val
            if section.margin_right is not None:
                val = safe_inches(section.margin_right)
                if val:
                    docx_section.right_margin = val
        
        for elem in section.elements:
            if isinstance(elem, Paragraph):
                self._write_paragraph(elem)
            elif isinstance(elem, Table):
                self._write_table(elem)
    
    def _write_paragraph(self, paragraph: Paragraph) -> None:
        """Write a paragraph."""
        if not self._doc:
            return
        
        docx_para = self._doc.add_paragraph()
        
        # Apply paragraph formatting
        para_format = docx_para.paragraph_format
        para_format.alignment = _get_docx_alignment(paragraph.alignment)
        
        if paragraph.space_before is not None:
            para_format.space_before = Pt(paragraph.space_before)
        if paragraph.space_after is not None:
            para_format.space_after = Pt(paragraph.space_after)
        if paragraph.left_indent is not None:
            para_format.left_indent = Inches(paragraph.left_indent)
        if paragraph.right_indent is not None:
            para_format.right_indent = Inches(paragraph.right_indent)
        if paragraph.first_line_indent is not None:
            para_format.first_line_indent = Inches(paragraph.first_line_indent)
        
        # Write runs
        for run in paragraph.runs:
            self._write_run(docx_para, run)
    
    def _write_run(self, docx_para, run: TextRun) -> None:
        """Write a text run."""
        if not run.text:
            return
        
        docx_run = docx_para.add_run(run.text)
        font = docx_run.font
        
        if run.font_name:
            font.name = run.font_name
        if run.font_size:
            font.size = Pt(run.font_size)
        
        font.bold = run.bold
        font.italic = run.italic
        font.underline = run.underline
        font.strike = run.strike
        font.superscript = run.superscript
        font.subscript = run.subscript
        
        if run.color:
            rgb = _hex_to_rgb(run.color)
            if rgb:
                font.color.rgb = rgb
    
    def _write_table(self, table: Table) -> None:
        """Write a table."""
        if not self._doc:
            return
        
        num_rows = table.num_rows
        num_cols = table.num_cols
        
        if num_rows == 0 or num_cols == 0:
            return
        
        docx_table = self._doc.add_table(rows=num_rows, cols=num_cols)
        
        # Apply table style if specified
        if table.style_name:
            try:
                docx_table.style = table.style_name
            except KeyError:
                # Style not found, use default
                docx_table.style = "Table Grid"
        else:
            docx_table.style = "Table Grid"
        
        # Write cells
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                if col_idx < num_cols:
                    docx_cell = docx_table.rows[row_idx].cells[col_idx]
                    self._write_cell(docx_cell, cell)
    
    def _write_cell(self, docx_cell, cell: TableCell) -> None:
        """Write a table cell."""
        # Clear default paragraph
        if docx_cell.paragraphs:
            docx_cell.paragraphs[0].clear()
        
        for i, para in enumerate(cell.paragraphs):
            if i == 0:
                # Use existing first paragraph
                docx_para = docx_cell.paragraphs[0]
                para_format = docx_para.paragraph_format
                para_format.alignment = _get_docx_alignment(para.alignment)
            else:
                docx_para = docx_cell.add_paragraph()
                para_format = docx_para.paragraph_format
                para_format.alignment = _get_docx_alignment(para.alignment)
            
            for run in para.runs:
                self._write_run(docx_para, run)


def write_docx(document: Document, output_path: Path | str) -> Path:
    """Convenience function to write a document to DOCX.
    
    Args:
        document: The document IR to write
        output_path: Path to save the DOCX file
        
    Returns:
        Path to the written file
    """
    writer = DocxWriter()
    return writer.write(document, output_path)


def write_docx_bytes(document: Document) -> bytes:
    """Convenience function to write a document to bytes.
    
    Args:
        document: The document IR to write
        
    Returns:
        DOCX file content as bytes
    """
    writer = DocxWriter()
    return writer.write_bytes(document)

