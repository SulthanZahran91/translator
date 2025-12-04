"""DOCX document parser that converts Word documents to IR.

This module parses .docx files using python-docx and converts them into
our Intermediate Representation (IR) format, preserving all formatting
metadata for later reconstruction.
"""

import base64
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from docx.document import Document as DocxDocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.table import Table as DocxTable
from docx.table import _Cell
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.text.run import Run as DocxRun

from backend.translation.ir import (
    Alignment,
    Document,
    DocumentStyle,
    Image,
    ImageType,
    Paragraph,
    Section,
    Table,
    TableCell,
    TableRow,
    TextRun,
)


def _pt_to_float(pt_value: object) -> float | None:
    """Convert Pt/EMU value to float points."""
    if pt_value is None:
        return None
    if hasattr(pt_value, "pt"):
        return float(pt_value.pt)
    if isinstance(pt_value, (int, float)):
        # Assume it's EMU, convert to points
        return float(pt_value) / 12700.0
    return None


def _inches_to_float(inches_value: object) -> float | None:
    """Convert Inches/EMU value to float inches.
    
    python-docx returns Length objects for dimension values. These have
    an `.inches` property. If we get a raw int/float, it's likely EMUs
    (English Metric Units), so we convert: 914400 EMUs = 1 inch.
    """
    if inches_value is None:
        return None
    # Check for Length objects first (they have .inches property)
    if hasattr(inches_value, "inches"):
        return float(inches_value.inches)
    # Raw int/float values are EMUs - convert to inches
    if isinstance(inches_value, (int, float)):
        # EMUs to inches: 914400 EMUs = 1 inch
        return float(inches_value) / 914400.0
    return None


def _rgb_to_hex(rgb: RGBColor | None) -> str | None:
    """Convert RGBColor to hex string."""
    if rgb is None:
        return None
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def _get_alignment(docx_alignment: WD_ALIGN_PARAGRAPH | None) -> Alignment:
    """Convert docx alignment to IR alignment."""
    if docx_alignment is None:
        return Alignment.LEFT

    alignment_map = {
        WD_ALIGN_PARAGRAPH.LEFT: Alignment.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER: Alignment.CENTER,
        WD_ALIGN_PARAGRAPH.RIGHT: Alignment.RIGHT,
        WD_ALIGN_PARAGRAPH.JUSTIFY: Alignment.JUSTIFY,
    }
    return alignment_map.get(docx_alignment, Alignment.LEFT)


class DocxParser:
    """Parser for converting DOCX files to IR format."""

    def __init__(self) -> None:
        self._doc: DocxDocumentType | None = None

    def parse(self, file_path: Path | str) -> Document:
        """Parse a DOCX file and return the IR Document.

        Args:
            file_path: Path to the .docx file

        Returns:
            Document IR representing the parsed content
        """
        file_path = Path(file_path)
        self._doc = DocxDocument(str(file_path))

        # Parse styles
        styles = self._parse_styles()

        # Parse sections
        sections = self._parse_sections()

        return Document(
            sections=sections,
            styles=styles,
            source_filename=file_path.name,
            source_format="docx",
        )

    def parse_bytes(self, content: bytes, filename: str = "document.docx") -> Document:
        """Parse DOCX content from bytes.

        Args:
            content: The DOCX file content as bytes
            filename: Original filename for metadata

        Returns:
            Document IR representing the parsed content
        """
        from io import BytesIO
        self._doc = DocxDocument(BytesIO(content))

        styles = self._parse_styles()
        sections = self._parse_sections()

        return Document(
            sections=sections,
            styles=styles,
            source_filename=filename,
            source_format="docx",
        )

    def _parse_styles(self) -> dict[str, DocumentStyle]:
        """Parse document styles."""
        styles: dict[str, DocumentStyle] = {}

        if self._doc is None:
            return styles

        for style in self._doc.styles:
            if style.type == 1:  # Paragraph style
                try:
                    font = style.font
                    para_format = style.paragraph_format

                    doc_style = DocumentStyle(
                        name=style.name,
                        base_style=style.base_style.name if style.base_style else None,
                        font_name=font.name if font else None,
                        font_size=_pt_to_float(font.size) if font else None,
                        bold=font.bold if font and font.bold else False,
                        italic=font.italic if font and font.italic else False,
                        alignment=_get_alignment(para_format.alignment) if para_format else Alignment.LEFT,
                        space_before=_pt_to_float(para_format.space_before) if para_format else None,
                        space_after=_pt_to_float(para_format.space_after) if para_format else None,
                    )
                    styles[style.name] = doc_style
                    styles[style.name] = doc_style
                except (AttributeError, ValueError):
                    continue

        return styles

    def _parse_sections(self) -> list[Section]:
        """Parse all sections in the document."""
        if self._doc is None:
            return []

        sections: list[Section] = []

        # Get section properties
        for docx_section in self._doc.sections:
            section = Section(
                page_width=_inches_to_float(docx_section.page_width),
                page_height=_inches_to_float(docx_section.page_height),
                margin_top=_inches_to_float(docx_section.top_margin),
                margin_bottom=_inches_to_float(docx_section.bottom_margin),
                margin_left=_inches_to_float(docx_section.left_margin),
                margin_right=_inches_to_float(docx_section.right_margin),
            )
            sections.append(section)

        # If no sections found, create a default one
        if not sections:
            sections.append(Section())

        # Parse body elements into the first section
        # (More complex section handling would require tracking section breaks)
        current_section = sections[0]
        current_section.elements = self._parse_body_elements()

        return sections

    def _parse_body_elements(self) -> list[Paragraph | Table | Image]:
        """Parse all body elements (paragraphs, tables, and images)."""
        if self._doc is None:
            return []

        elements: list[Paragraph | Table | Image] = []

        # python-docx exposes body elements through the document body
        body = self._doc.element.body

        for child in body:
            tag = child.tag

            if tag.endswith("}p"):  # Paragraph
                # Find the corresponding paragraph object
                for para in self._doc.paragraphs:
                    if para._element is child:
                        # Check for images in this paragraph and extract them
                        images = self._extract_images_from_paragraph(para)
                        elements.extend(images)
                        
                        # Parse the paragraph itself (may have text alongside images)
                        parsed_para = self._parse_paragraph(para)
                        
                        # TRS 1.3 Image Guard Logic: If paragraph contains images,
                        # mark it to skip translation to prevent image anchor deletion
                        if images:
                            parsed_para.skip_translation = True
                        
                        # Only add if it has content or non-empty runs
                        if parsed_para.runs and any(r.text for r in parsed_para.runs):
                            elements.append(parsed_para)
                        elif not images:
                            # Empty paragraph with no images - keep for spacing
                            elements.append(parsed_para)
                        break
            elif tag.endswith("}tbl"):  # Table
                for table in self._doc.tables:
                    if table._element is child:
                        elements.append(self._parse_table(table))
                        break

        return elements

    def _extract_images_from_paragraph(self, docx_para: DocxParagraph) -> list[Image]:
        """Extract all images from a paragraph's runs."""
        images: list[Image] = []
        
        for run in docx_para.runs:
            # Check for inline images (drawings)
            run_images = self._extract_images_from_run(run)
            images.extend(run_images)
        
        return images
    
    def _extract_images_from_run(self, docx_run: DocxRun) -> list[Image]:
        """Extract images from a single run element."""
        images: list[Image] = []
        
        if self._doc is None:
            return images
        
        run_element = docx_run._r
        
        # Look for drawing elements (inline images)
        # Namespace for drawings
        drawing_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
        
        for drawing in run_element.iter(drawing_ns):
            image = self._parse_drawing_element(drawing, ImageType.INLINE)
            if image:
                images.append(image)
        
        return images
    
    def _parse_drawing_element(self, drawing_elem, image_type: ImageType) -> Image | None:
        """Parse a drawing element and extract image data."""
        if self._doc is None:
            return None
        
        # Namespaces used in OOXML
        ns = {
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
        }
        
        try:
            # Find the blip element which contains the image reference
            blip = drawing_elem.find('.//a:blip', ns)
            if blip is None:
                return None
            
            # Get the relationship ID
            embed_attr = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
            rel_id = blip.get(embed_attr)
            if not rel_id:
                return None
            
            # Get the image part from the document's relationships
            try:
                image_part = self._doc.part.related_parts[rel_id]
            except KeyError:
                return None
            
            # Get image data
            image_bytes = image_part.blob
            image_data = base64.b64encode(image_bytes).decode('utf-8')
            
            # Determine image format from content type
            content_type = image_part.content_type
            format_map = {
                'image/png': 'png',
                'image/jpeg': 'jpeg',
                'image/jpg': 'jpeg',
                'image/gif': 'gif',
                'image/bmp': 'bmp',
                'image/tiff': 'tiff',
                'image/x-emf': 'emf',
                'image/x-wmf': 'wmf',
            }
            image_format = format_map.get(content_type, 'png')
            
            # Try to extract dimensions
            width = None
            height = None
            
            # Look for extent element (dimensions in EMUs)
            extent = drawing_elem.find('.//wp:extent', ns)
            if extent is not None:
                cx = extent.get('cx')
                cy = extent.get('cy')
                if cx:
                    # Convert EMUs to inches (914400 EMUs = 1 inch)
                    width = int(cx) / 914400.0
                if cy:
                    height = int(cy) / 914400.0
            
            # Try to get alt text
            alt_text = None
            doc_pr = drawing_elem.find('.//wp:docPr', ns)
            if doc_pr is not None:
                alt_text = doc_pr.get('descr')
            
            # Try to get positioning for floating images
            position_x = None
            position_y = None
            
            if image_type == ImageType.FLOATING:
                # Look for anchor positioning
                pos_h = drawing_elem.find('.//wp:positionH', ns)
                pos_v = drawing_elem.find('.//wp:positionV', ns)
                if pos_h is not None:
                    pos_offset = pos_h.find('.//wp:posOffset', ns)
                    if pos_offset is not None and pos_offset.text:
                        position_x = int(pos_offset.text) / 914400.0
                if pos_v is not None:
                    pos_offset = pos_v.find('.//wp:posOffset', ns)
                    if pos_offset is not None and pos_offset.text:
                        position_y = int(pos_offset.text) / 914400.0
            
            return Image(
                data=image_data,
                format=image_format,
                width=width,
                height=height,
                image_type=image_type,
                position_x=position_x,
                position_y=position_y,
                alt_text=alt_text,
                rel_id=rel_id,
            )
            
        except Exception:
            # If anything goes wrong, skip this image
            return None

    def _parse_paragraph(self, docx_para: DocxParagraph) -> Paragraph:
        """Parse a single paragraph."""
        runs = [self._parse_run(run) for run in docx_para.runs]

        para_format = docx_para.paragraph_format

        return Paragraph(
            runs=runs,
            style_name=docx_para.style.name if docx_para.style else None,
            alignment=_get_alignment(para_format.alignment),
            space_before=_pt_to_float(para_format.space_before),
            space_after=_pt_to_float(para_format.space_after),
            line_spacing=_pt_to_float(para_format.line_spacing),
            left_indent=_inches_to_float(para_format.left_indent),
            right_indent=_inches_to_float(para_format.right_indent),
            first_line_indent=_inches_to_float(para_format.first_line_indent),
        )

    def _parse_run(self, docx_run: DocxRun) -> TextRun:
        """Parse a single text run."""
        font = docx_run.font

        return TextRun(
            text=docx_run.text,
            font_name=font.name,
            font_size=_pt_to_float(font.size),
            bold=font.bold if font.bold is not None else False,
            italic=font.italic if font.italic is not None else False,
            underline=font.underline is not None and font.underline is not False,
            strike=font.strike if font.strike is not None else False,
            color=_rgb_to_hex(font.color.rgb) if font.color and font.color.rgb else None,
            highlight_color=str(font.highlight_color) if font.highlight_color else None,
            superscript=font.superscript if font.superscript is not None else False,
            subscript=font.subscript if font.subscript is not None else False,
        )

    def _parse_table(self, docx_table: DocxTable) -> Table:
        """Parse a table."""
        rows: list[TableRow] = []
        col_widths: list[float] = []

        # Try to get column widths
        try:
            tbl = docx_table._tbl
            tbl_grid = tbl.tblGrid
            if tbl_grid is not None:
                for grid_col in tbl_grid.gridCol_lst:
                    width = grid_col.get(qn("w:w"))
                    if width:
                        # Convert twips to inches (1 inch = 1440 twips)
                        col_widths.append(int(width) / 1440.0)
        except (AttributeError, ValueError):
            pass

        for docx_row in docx_table.rows:
            cells: list[TableCell] = []

            for docx_cell in docx_row.cells:
                cell = self._parse_cell(docx_cell)
                cells.append(cell)

            row = TableRow(
                cells=cells,
                height=_inches_to_float(docx_row.height),
            )
            rows.append(row)

        return Table(
            rows=rows,
            col_widths=col_widths,
            style_name=docx_table.style.name if docx_table.style else None,
        )

    def _parse_cell(self, docx_cell: _Cell) -> TableCell:
        """Parse a table cell."""
        paragraphs = [self._parse_paragraph(p) for p in docx_cell.paragraphs]

        # Get cell properties
        tc = docx_cell._tc
        tc_pr = tc.tcPr

        # Get vertical merge info for row span
        row_span = 1
        col_span = 1

        if tc_pr is not None:
            # Column span
            grid_span = tc_pr.find(qn("w:gridSpan"))
            if grid_span is not None:
                val = grid_span.get(qn("w:val"))
                if val:
                    col_span = int(val)

            # Vertical alignment
            v_align = tc_pr.find(qn("w:vAlign"))
            v_alignment = None
            if v_align is not None:
                v_alignment = v_align.get(qn("w:val"))
        else:
            v_alignment = None

        return TableCell(
            paragraphs=paragraphs,
            width=_inches_to_float(docx_cell.width),
            row_span=row_span,
            col_span=col_span,
            vertical_alignment=v_alignment,
        )


def parse_docx(file_path: Path | str) -> Document:
    """Convenience function to parse a DOCX file.

    Args:
        file_path: Path to the .docx file

    Returns:
        Document IR representing the parsed content
    """
    parser = DocxParser()
    return parser.parse(file_path)


def parse_docx_bytes(content: bytes, filename: str = "document.docx") -> Document:
    """Convenience function to parse DOCX from bytes.

    Args:
        content: The DOCX file content as bytes
        filename: Original filename for metadata

    Returns:
        Document IR representing the parsed content
    """
    parser = DocxParser()
    return parser.parse_bytes(content, filename)

