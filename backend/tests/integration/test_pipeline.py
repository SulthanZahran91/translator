"""Integration tests for the full translation pipeline."""

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document as DocxDocument

from backend.translation.ingestion.docx_parser import parse_docx_bytes
from backend.translation.chunking.chunker import chunk_document, ChunkingConfig
from backend.translation.orchestrator import MockTranslationOrchestrator
from backend.translation.reconstruction.reconstructor import reconstruct_document
from backend.translation.export.docx_writer import write_docx_bytes
from backend.translation.glossary.manager import GlossaryManager
from backend.translation.ir import Document


def create_test_document_bytes() -> bytes:
    """Create a test DOCX document with Korean content."""
    doc = DocxDocument()
    
    # Add heading
    doc.add_heading("테스트 문서", level=1)
    
    # Add paragraphs
    doc.add_paragraph("안녕하세요, 이것은 테스트입니다.")
    doc.add_paragraph("번역 시스템을 테스트하고 있습니다.")
    
    # Add a simple table
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "항목"
    table.rows[0].cells[1].text = "값"
    table.rows[1].cells[0].text = "테스트"
    table.rows[1].cells[1].text = "성공"
    
    # Add conclusion
    doc.add_paragraph("감사합니다.")
    
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


class TestFullPipeline:
    """Integration tests for the complete translation pipeline."""
    
    @pytest.mark.asyncio
    async def test_ingest_chunk_translate_reconstruct_export(self):
        """Test the full pipeline: ingest → chunk → translate → reconstruct → export."""
        # 1. Create test document
        input_bytes = create_test_document_bytes()
        
        # 2. Parse document to IR
        document = parse_docx_bytes(input_bytes, "test.docx")
        assert isinstance(document, Document)
        assert len(document.all_paragraphs) > 0
        
        # 3. Chunk document
        config = ChunkingConfig(max_tokens_per_unit=10000)
        units = chunk_document(document, config)
        assert len(units) >= 1
        
        # 4. Translate units (using mock)
        mock_responses = {}
        for i, unit in enumerate(units):
            # Create mock translations that preserve structure
            translated = unit.source_text
            # Simple mock translation: add [EN] prefix to content
            translated = translated.replace("테스트 문서", "Test Document")
            translated = translated.replace("안녕하세요", "Hello")
            translated = translated.replace("번역 시스템을 테스트하고 있습니다", "Testing the translation system")
            translated = translated.replace("감사합니다", "Thank you")
            translated = translated.replace("항목", "Item")
            translated = translated.replace("값", "Value")
            translated = translated.replace("테스트", "Test")
            translated = translated.replace("성공", "Success")
            mock_responses[i] = translated
        
        orchestrator = MockTranslationOrchestrator(mock_responses=mock_responses)
        translated_units = await orchestrator.translate_units(units)
        
        assert len(translated_units) == len(units)
        for unit in translated_units:
            assert unit.translated_text is not None
        
        # 5. Reconstruct document
        result = reconstruct_document(document, translated_units)
        assert result.elements_updated > 0
        
        # 6. Export to DOCX
        output_bytes = write_docx_bytes(document)
        assert len(output_bytes) > 0
        
        # 7. Verify output can be read
        output_doc = DocxDocument(BytesIO(output_bytes))
        all_text = " ".join(p.text for p in output_doc.paragraphs)
        
        # Should contain translated content
        assert "Test Document" in all_text or "Hello" in all_text or "Test" in all_text
    
    @pytest.mark.asyncio
    async def test_glossary_consistency_across_units(self):
        """Test that glossary terms remain consistent across translation units."""
        # Create document with repeated terms
        doc = DocxDocument()
        doc.add_paragraph("이 문서에서 기술용어를 사용합니다.")
        doc.add_paragraph("기술용어는 중요합니다.")
        doc.add_paragraph("다시 기술용어를 언급합니다.")
        
        buffer = BytesIO()
        doc.save(buffer)
        input_bytes = buffer.getvalue()
        
        # Parse and chunk
        document = parse_docx_bytes(input_bytes)
        
        # Use small chunk size to create multiple units
        config = ChunkingConfig(max_tokens_per_unit=100)
        units = chunk_document(document, config)
        
        # Create mock responses that include glossary tags in first unit
        # and use the term in subsequent units
        mock_responses = {}
        for i, unit in enumerate(units):
            if i == 0:
                # First unit extracts the term
                mock_responses[i] = unit.source_text.replace(
                    "기술용어",
                    "technical term <glossary>기술용어|technical term</glossary>"
                )
            else:
                # Subsequent units should use the same translation
                mock_responses[i] = unit.source_text.replace("기술용어", "technical term")
        
        glossary_manager = GlossaryManager()
        orchestrator = MockTranslationOrchestrator(
            mock_responses=mock_responses,
            glossary_manager=glossary_manager,
        )
        
        await orchestrator.translate_units(units)
        
        # Check that the term was extracted and stored
        job_glossary = glossary_manager.get_job_glossary()
        assert job_glossary.has_term("기술용어")
        assert job_glossary.get_term("기술용어").target_term == "technical term"
    
    @pytest.mark.asyncio
    async def test_checkpoint_and_resume(self):
        """Test that translation can be paused and resumed."""
        doc = DocxDocument()
        for i in range(10):
            doc.add_paragraph(f"문단 {i}")
        
        buffer = BytesIO()
        doc.save(buffer)
        input_bytes = buffer.getvalue()
        
        document = parse_docx_bytes(input_bytes)
        config = ChunkingConfig(max_tokens_per_unit=100)
        units = chunk_document(document, config)
        
        # First run: translate first half
        orchestrator1 = MockTranslationOrchestrator()
        
        # Pause after translating some units
        mid_point = len(units) // 2
        partial_units = units[:mid_point]
        for i, u in enumerate(partial_units):
            u.sequence_number = i
        
        result1 = await orchestrator1.translate_units(partial_units)
        assert len(result1) == mid_point
        
        # Simulate saving and restoring state
        # In real usage, this would go through CheckpointManager
        saved_glossary = orchestrator1.glossary_manager.to_dict()
        
        # Second run: continue from where we left off
        orchestrator2 = MockTranslationOrchestrator(
            glossary_manager=GlossaryManager.from_dict(saved_glossary),
        )
        
        # Copy translated text from first run
        for i in range(mid_point):
            units[i].translated_text = result1[i].translated_text
        
        result2 = await orchestrator2.translate_units(units, start_from=mid_point)
        
        # All units should now be translated
        assert len(result2) == len(units)
        for unit in result2:
            assert unit.translated_text is not None


class TestPipelineWithTables:
    """Integration tests focusing on table handling."""
    
    @pytest.mark.asyncio
    async def test_table_translation_preserves_structure(self):
        """Test that table structure is preserved through translation."""
        doc = DocxDocument()
        doc.add_paragraph("테이블 소개:")
        
        table = doc.add_table(rows=3, cols=3)
        for i in range(3):
            for j in range(3):
                table.rows[i].cells[j].text = f"셀{i}{j}"
        
        doc.add_paragraph("테이블 끝.")
        
        buffer = BytesIO()
        doc.save(buffer)
        input_bytes = buffer.getvalue()
        
        # Parse
        document = parse_docx_bytes(input_bytes)
        
        # Verify table was parsed
        tables = document.all_tables
        assert len(tables) == 1
        assert tables[0].num_rows == 3
        assert tables[0].num_cols == 3
        
        # Chunk and translate
        units = chunk_document(document)
        
        mock_responses = {}
        for i, unit in enumerate(units):
            # Simple mock: replace Korean with English
            translated = unit.source_text
            for row in range(3):
                for col in range(3):
                    translated = translated.replace(f"셀{row}{col}", f"Cell{row}{col}")
            translated = translated.replace("테이블 소개", "Table introduction")
            translated = translated.replace("테이블 끝", "Table end")
            mock_responses[i] = translated
        
        orchestrator = MockTranslationOrchestrator(mock_responses=mock_responses)
        translated_units = await orchestrator.translate_units(units)
        
        # Reconstruct
        result = reconstruct_document(document, translated_units)
        
        # Verify table structure is preserved
        tables = document.all_tables
        assert len(tables) == 1
        assert tables[0].num_rows == 3
        assert tables[0].num_cols == 3


class TestPipelineErrorHandling:
    """Tests for error handling in the pipeline."""
    
    def test_empty_document(self):
        """Test handling of empty document."""
        doc = DocxDocument()
        
        buffer = BytesIO()
        doc.save(buffer)
        input_bytes = buffer.getvalue()
        
        document = parse_docx_bytes(input_bytes)
        units = chunk_document(document)
        
        # Should handle empty document gracefully
        assert len(units) == 0
    
    def test_document_with_only_whitespace(self):
        """Test handling of document with only whitespace."""
        doc = DocxDocument()
        doc.add_paragraph("   ")
        doc.add_paragraph("\n\n")
        
        buffer = BytesIO()
        doc.save(buffer)
        input_bytes = buffer.getvalue()
        
        document = parse_docx_bytes(input_bytes)
        units = chunk_document(document)
        
        # Whitespace-only paragraphs should be skipped
        assert len(units) == 0

